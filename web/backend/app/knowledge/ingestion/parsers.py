"""文档解析。"""
from __future__ import annotations

import json
from html.parser import HTMLParser
from pathlib import Path

from app.services.case_kb import build_kb_search_blob
from app.models import TestCase


class _HtmlTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        t = data.strip()
        if t:
            self.parts.append(t)


def parse_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def parse_markdown(path: Path) -> str:
    return parse_text_file(path)


def parse_json_file(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        data = json.loads(raw)
        return json.dumps(data, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        return raw


def parse_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_pdf(path: Path) -> str:
    import fitz

    doc = fitz.open(str(path))
    parts: list[str] = []
    for page in doc:
        t = page.get_text().strip()
        if t:
            parts.append(t)
    doc.close()
    return "\n\n".join(parts)


def parse_html(path: Path) -> str:
    raw = parse_text_file(path)
    parser = _HtmlTextExtractor()
    try:
        parser.feed(raw)
    except Exception:
        return raw
    return "\n".join(parser.parts) if parser.parts else raw


def parse_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in wb.worksheets:
            parts.append(f"## {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
    finally:
        wb.close()
    return "\n".join(parts)


def parse_xls(path: Path) -> str:
    import xlrd

    book = xlrd.open_workbook(str(path))
    parts: list[str] = []
    for sheet in book.sheets():
        parts.append(f"## {sheet.name}")
        for rx in range(sheet.nrows):
            cells = [
                str(sheet.cell_value(rx, cx)).strip()
                for cx in range(sheet.ncols)
                if str(sheet.cell_value(rx, cx)).strip()
            ]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_case_from_db(case: TestCase) -> str:
    return build_kb_search_blob(case)


def parse_structured_json(structured_json: str, doc_type: str) -> str:
    try:
        data = json.loads(structured_json or "{}")
    except json.JSONDecodeError:
        return structured_json or ""
    if doc_type == "page_model":
        name = data.get("page_name") or data.get("name") or ""
        desc = data.get("description") or ""
        elements = data.get("elements") or []
        lines = [f"页面：{name}", desc]
        for el in elements if isinstance(elements, list) else []:
            if isinstance(el, dict):
                lines.append(
                    f"控件：{el.get('name','')} 类型={el.get('type','')} 定位={el.get('locator','')}"
                )
        return "\n".join(x for x in lines if x)
    if doc_type == "ui_element":
        return json.dumps(data, ensure_ascii=False, indent=2)
    return json.dumps(data, ensure_ascii=False, indent=2)


def parse_feature_tree_json(tree_json: str) -> str:
    try:
        tree = json.loads(tree_json or "{}")
    except json.JSONDecodeError:
        return tree_json or ""
    features = tree.get("features") or []
    lines: list[str] = [f"应用：{tree.get('app_name','')}", f"Bundle：{tree.get('bundle_id','')}"]
    for f in features if isinstance(features, list) else []:
        if not isinstance(f, dict):
            continue
        path = " > ".join(str(x) for x in (f.get("path") or []) if x)
        lines.append(
            f"功能：{f.get('name','')} 区域={f.get('region','')} 路径={path} 屏幕={f.get('screen_title','')}"
        )
    return "\n".join(lines)


def parse_document_content(
    *,
    file_path: str | None,
    structured_json: str,
    doc_type: str,
    case: TestCase | None = None,
    feature_tree_json: str | None = None,
) -> str:
    if case is not None:
        return parse_case_from_db(case)
    if feature_tree_json:
        return parse_feature_tree_json(feature_tree_json)
    if structured_json and structured_json.strip() not in ("", "{}"):
        parsed = parse_structured_json(structured_json, doc_type)
        if parsed.strip():
            return parsed
    if not file_path:
        return ""
    path = Path(file_path)
    if not path.is_file():
        return ""
    suffix = path.suffix.lower()
    if suffix in (".md", ".txt", ".markdown", ".mdx"):
        return parse_markdown(path)
    if suffix in (".html", ".htm"):
        return parse_html(path)
    if suffix == ".csv":
        return parse_text_file(path)
    if suffix == ".json":
        return parse_json_file(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".xlsx":
        return parse_xlsx(path)
    if suffix == ".xls":
        return parse_xls(path)
    return parse_text_file(path)
